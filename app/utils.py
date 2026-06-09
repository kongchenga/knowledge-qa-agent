"""Enhanced document parsing: PDF (PyMuPDF), DOCX (python-docx), XLSX (openpyxl).

Returns a ParseResult with markdown text and metadata (images, tables, headings).
"""
import io
import uuid
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

from app.config import settings

IMAGES_DIR = settings.resolved_knowledge_dir / "images"


@dataclass
class ParseResult:
    text: str = ""
    images: list[str] = field(default_factory=list)
    tables: int = 0
    headings: list[str] = field(default_factory=list)


# ─── public API ──────────────────────────────────────────────────────────────

def extract_text_from_file(file_path: Path) -> Optional[str]:
    suffix = file_path.suffix.lower()
    result = _route_extract(file_path, suffix)
    return result.text if result else None


def extract_text_from_bytes(content: bytes, filename: str) -> Optional[str]:
    suffix = Path(filename).suffix.lower()
    if suffix in _PLAIN_EXTENSIONS:
        return content.decode("utf-8")
    result = _route_extract_bytes(content, suffix, filename)
    return result.text if result else None


def parse_document(content: bytes, filename: str) -> Optional[ParseResult]:
    """Full parse returning structured ParseResult with text + metadata."""
    suffix = Path(filename).suffix.lower()
    if suffix in _PLAIN_EXTENSIONS:
        return ParseResult(text=content.decode("utf-8"))
    return _route_extract_bytes(content, suffix, filename)


_PLAIN_EXTENSIONS = {".txt", ".md", ".py", ".json", ".yaml", ".yml",
                     ".csv", ".xml", ".html", ".css", ".js"}


def _route_extract(path: Path, suffix: str) -> Optional[ParseResult]:
    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".xlsx":
        return _extract_xlsx(path)
    return None


def _route_extract_bytes(content: bytes, suffix: str, filename: str) -> Optional[ParseResult]:
    if suffix == ".pdf":
        return _extract_pdf_bytes(content)
    elif suffix == ".docx":
        return _extract_docx_bytes(content)
    elif suffix == ".xlsx":
        return _extract_xlsx_bytes(content)
    return None


# ─── PDF (PyMuPDF) ───────────────────────────────────────────────────────────

def _extract_pdf(path: Path) -> Optional[ParseResult]:
    import fitz
    return _parse_pdf_doc(fitz.open(str(path)))


def _extract_pdf_bytes(content: bytes) -> Optional[ParseResult]:
    import fitz
    return _parse_pdf_doc(fitz.open(stream=content, filetype="pdf"))


def _parse_pdf_doc(doc) -> Optional[ParseResult]:
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    result = ParseResult()
    pages_md = []

    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        page_lines = []

        for block in blocks:
            if block["type"] == 0:  # text block
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    prefix = _detect_heading(spans[0])
                    text = "".join(s.get("text", "") for s in spans).strip()
                    if text:
                        page_lines.append(f"{prefix}{text}\n")
            elif block["type"] == 1:  # image block
                try:
                    img_refs = _extract_image_by_xref(doc, page_num)
                    for ir in img_refs:
                        result.images.append(ir)
                        page_lines.append(f"![image]({ir})\n\n")
                except Exception:
                    pass

        pages_md.append("".join(page_lines))

    tables = _extract_pdf_tables(doc)
    table_md_sections = []
    for i, tbl in enumerate(tables):
        header = tbl.get("header", [])
        rows = tbl.get("rows", [])
        md = _format_markdown_table(header, rows)
        table_md_sections.append(f"### 表格 {i + 1}\n\n{md}\n")
    result.tables = len(tables)

    result.text = "\n\n".join(pages_md + table_md_sections)
    return result


def _detect_heading(span: dict) -> str:
    size = span.get("size", 12)
    flags = span.get("flags", 0)
    is_bold = bool(flags & 2)
    font = span.get("font", "").lower()

    if "bold" in font or is_bold:
        if size >= 24:
            return "# "
        elif size >= 18:
            return "## "
        elif size >= 14:
            return "### "
    elif size >= 24:
        return "## "
    elif size >= 18:
        return "### "
    return ""


def _extract_image_by_xref(doc, page_num: int) -> list[str]:
    refs = []
    relpath = "images"
    for img_info in doc.get_page_images(page_num):
        xref = img_info[0]
        img_bytes = doc.extract_image(xref)
        ext = img_bytes.get("ext", "png")
        fname = f"{uuid.uuid4().hex}.{ext}"
        fpath = IMAGES_DIR / fname
        fpath.write_bytes(img_bytes["image"])
        refs.append(f"{relpath}/{fname}")
    return refs


def _extract_pdf_tables(doc) -> list[dict]:
    try:
        tables = []
        for page in doc:
            found = page.find_tables()
            for t in found:
                data = t.extract()
                if data:
                    tables.append({"header": data[0] if len(data) > 1 else [],
                                   "rows": data[1:]})
        return tables
    except Exception:
        return []


# ─── DOCX (python-docx) ──────────────────────────────────────────────────────

def _extract_docx(path: Path) -> Optional[ParseResult]:
    from docx import Document as DocxDoc
    return _parse_docx_doc(DocxDoc(str(path)))


def _extract_docx_bytes(content: bytes) -> Optional[ParseResult]:
    from docx import Document as DocxDoc
    return _parse_docx_doc(DocxDoc(io.BytesIO(content)))


def _parse_docx_doc(doc) -> Optional[ParseResult]:
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)

    result = ParseResult()
    md_lines = []
    tables = 0

    _docx_extract_part_images(doc, result.images)

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            for para in doc.paragraphs:
                if para._element is element:
                    md_lines.append(_format_docx_paragraph(para))
                    break
        elif tag == "tbl":
            tables += 1
            # find corresponding python-docx Table
            for tbl in doc.tables:
                if tbl._element is element:
                    md_lines.append(_format_docx_table(tbl, tables))
                    break

    result.text = "\n".join(md_lines)
    result.tables = tables
    return result


def _format_docx_paragraph(para) -> str:
    """Convert a paragraph to Markdown, preserving bold/italic/headings."""
    style = para.style.name.lower() if para.style else ""
    text_parts = []
    for run in para.runs:
        txt = run.text
        if not txt:
            continue
        if run.bold and run.italic:
            txt = f"***{txt}***"
        elif run.bold:
            txt = f"**{txt}**"
        elif run.italic:
            txt = f"*{txt}*"
        text_parts.append(txt)

    line = "".join(text_parts)

    if "heading 1" in style or "title" in style:
        return f"# {line}\n"
    elif "heading 2" in style:
        return f"## {line}\n"
    elif "heading 3" in style:
        return f"### {line}\n"
    elif "heading" in style:
        return f"#### {line}\n"
    return f"{line}\n"


def _format_docx_table(tbl, idx: int) -> str:
    """Convert a docx table to Markdown."""
    rows = []
    for row in tbl.rows:
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = rows[0] if rows else []
    data = rows[1:] if len(rows) > 1 else []
    return f"### 表格 {idx}\n\n{_format_markdown_table(header, data)}\n"


def _docx_extract_part_images(doc, images_list: list[str]):
    """Extract images embedded in docx relationships."""
    relpath = "images"
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                ext = rel.target_part.partname.split(".")[-1]
                if ext not in ("png", "jpg", "jpeg", "gif", "bmp", "webp", "svg"):
                    ext = "png"
                fname = f"{uuid.uuid4().hex}.{ext}"
                fpath = IMAGES_DIR / fname
                fpath.write_bytes(img_bytes)
                md_img = f"![image]({relpath}/{fname})"
                images_list.append(md_img)
            except Exception:
                pass


# ─── XLSX (openpyxl) ─────────────────────────────────────────────────────────

def _extract_xlsx(path: Path) -> Optional[ParseResult]:
    from openpyxl import load_workbook
    return _parse_xlsx_workbook(load_workbook(str(path), data_only=True))


def _extract_xlsx_bytes(content: bytes) -> Optional[ParseResult]:
    from openpyxl import load_workbook
    return _parse_xlsx_workbook(load_workbook(io.BytesIO(content), data_only=True))


def _parse_xlsx_workbook(wb) -> Optional[ParseResult]:
    result = ParseResult()
    sections = []
    table_count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(c) if c is not None else "" for c in row])
        if not rows:
            continue
        table_count += 1
        sections.append(f"## {sheet_name}\n\n{_format_markdown_table(rows[0], rows[1:])}")

    result.text = "\n\n".join(sections)
    result.tables = table_count
    return result


# ─── helpers ─────────────────────────────────────────────────────────────────

def _format_markdown_table(header: list[str | None], rows: list[list[str | None]]) -> str:
    if not header:
        return ""
    header = [(h or "").replace("\n", " ").strip() for h in header]
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows:
        r = [(c or "").replace("\n", " ").strip() for c in row]
        if len(r) < len(header):
            r += [""] * (len(header) - len(r))
        lines.append("| " + " | ".join(r[:len(header)]) + " |")
    return "\n".join(lines)
